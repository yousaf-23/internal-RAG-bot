"""
Document Processor Module
=========================
Purpose: Extract text from various document types (PDF, Word, Excel)
Author: RAG System Development
Date: 2024

This module handles:
1. Text extraction from different file formats
2. Text chunking for vector embedding
3. Metadata extraction (page numbers, word counts)
"""

# ============================================================================
# IMPORTS SECTION
# ============================================================================
# We need different libraries for different file types

import os
import re
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime

# For PDF processing
try:
    from pypdf import PdfReader
    PDF_AVAILABLE = True
    print("[DocumentProcessor] ✅ PDF support available (pypdf)")
except ImportError:
    PDF_AVAILABLE = False
    print("[DocumentProcessor] ❌ PDF support not available - install pypdf")

# For Word document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    print("[DocumentProcessor] ✅ Word support available (python-docx)")
except ImportError:
    DOCX_AVAILABLE = False
    print("[DocumentProcessor] ❌ Word support not available - install python-docx")

# For Excel processing
try:
    import openpyxl
    EXCEL_AVAILABLE = True
    print("[DocumentProcessor] ✅ Excel support available (openpyxl)")
except ImportError:
    EXCEL_AVAILABLE = False
    print("[DocumentProcessor] ❌ Excel support not available - install openpyxl")

# For text splitting (important for creating chunks)
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    LANGCHAIN_AVAILABLE = True
    print("[DocumentProcessor] ✅ LangChain text splitter available")
except ImportError:
    LANGCHAIN_AVAILABLE = False
    print("[DocumentProcessor] ⚠️  LangChain not available - using basic splitter")

# Import our configuration
from app.config import settings

# ============================================================================
# DOCUMENT PROCESSOR CLASS
# ============================================================================

class DocumentProcessor:
    """
    Main class for processing documents.
    
    This class handles:
    - Text extraction from various file formats
    - Text chunking for vector embeddings
    - Metadata extraction
    """
    
    def __init__(self):
        """
        Initialize the document processor with configuration settings.
        """
        print("\n[DocumentProcessor] Initializing...")
        
        # Configuration from settings
        self.chunk_size = settings.chunk_size  # Default: 1000 tokens
        self.chunk_overlap = settings.chunk_overlap  # Default: 200 tokens
        
        # Supported file extensions
        self.supported_extensions = {
            'pdf': self.process_pdf,
            'docx': self.process_docx,
            'doc': self.process_docx,  # Treat .doc same as .docx
            'xlsx': self.process_excel,
            'xls': self.process_excel,  # Treat .xls same as .xlsx
            'txt': self.process_text  # Plain text files
        }
        
        # Initialize text splitter if LangChain is available
        if LANGCHAIN_AVAILABLE:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                length_function=len,  # Use character count
                separators=["\n\n", "\n", ".", "!", "?", ";", ",", " ", ""],
                is_separator_regex=False
            )
            print(f"[DocumentProcessor] Text splitter configured:")
            print(f"  - Chunk size: {self.chunk_size} characters")
            print(f"  - Overlap: {self.chunk_overlap} characters")
        else:
            self.text_splitter = None
            print("[DocumentProcessor] Using basic text splitter")
        
        print(f"[DocumentProcessor] Supported formats: {list(self.supported_extensions.keys())}")
        print("[DocumentProcessor] ✅ Initialization complete\n")
    
    # ========================================================================
    # MAIN PROCESSING METHOD
    # ========================================================================
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Main method to process any document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing:
            - success: Whether processing succeeded
            - text: Full extracted text
            - chunks: List of text chunks
            - metadata: Document metadata (pages, word count, etc.)
            - error: Error message if failed
        """
        print(f"\n{'='*60}")
        print(f"[DocumentProcessor] Starting processing")
        print(f"  File: {file_path}")
        print(f"{'='*60}")
        
        # Step 1: Validate file exists
        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            return {
                'success': False,
                'error': error_msg
            }
        
        # Step 2: Get file information
        file_path = Path(file_path)
        file_size = file_path.stat().st_size
        file_extension = file_path.suffix.lower().strip('.')
        
        print(f"[DocumentProcessor] File info:")
        print(f"  - Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)")
        print(f"  - Extension: .{file_extension}")
        
        # Step 3: Check if file type is supported
        if file_extension not in self.supported_extensions:
            error_msg = f"Unsupported file type: .{file_extension}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            return {
                'success': False,
                'error': error_msg,
                'supported_types': list(self.supported_extensions.keys())
            }
        
        # Step 4: Call the appropriate processor
        try:
            print(f"[DocumentProcessor] Processing as {file_extension.upper()}...")
            processor_function = self.supported_extensions[file_extension]
            result = processor_function(str(file_path))
            
            # Step 5: If successful, create chunks
            if result['success'] and 'text' in result:
                print(f"[DocumentProcessor] Creating chunks from extracted text...")
                chunks = self.create_chunks(result['text'])
                result['chunks'] = chunks
                print(f"[DocumentProcessor] ✅ Created {len(chunks)} chunks")
            
            return result
            
        except Exception as e:
            error_msg = f"Processing failed: {str(e)}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            
            # Print detailed error for debugging
            import traceback
            print("\n[DocumentProcessor] Detailed error:")
            traceback.print_exc()
            
            return {
                'success': False,
                'error': error_msg
            }
    
    # ========================================================================
    # PDF PROCESSING
    # ========================================================================
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF files.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        print("[DocumentProcessor] Extracting text from PDF...")
        
        if not PDF_AVAILABLE:
            return {
                'success': False,
                'error': 'PDF support not available. Install pypdf.'
            }
        
        try:
            # Open and read the PDF
            reader = PdfReader(file_path)
            num_pages = len(reader.pages)
            print(f"[DocumentProcessor] PDF has {num_pages} pages")
            
            # Extract text from each page
            full_text = ""
            page_texts = []
            word_count = 0
            
            for page_num, page in enumerate(reader.pages, 1):
                # Extract text from this page
                page_text = page.extract_text()
                
                # Clean up the text (remove extra whitespace)
                page_text = re.sub(r'\s+', ' ', page_text).strip()
                
                if page_text:
                    page_texts.append({
                        'page': page_num,
                        'text': page_text,
                        'word_count': len(page_text.split())
                    })
                    
                    full_text += f"\n\n[Page {page_num}]\n{page_text}"
                    word_count += len(page_text.split())
                    
                    # Debug: Show progress for large PDFs
                    if page_num % 10 == 0:
                        print(f"  Processed {page_num}/{num_pages} pages...")
            
            print(f"[DocumentProcessor] ✅ Extracted text from {len(page_texts)} pages")
            print(f"[DocumentProcessor] Total words: {word_count:,}")
            
            return {
                'success': True,
                'text': full_text.strip(),
                'metadata': {
                    'page_count': num_pages,
                    'word_count': word_count,
                    'pages_with_text': len(page_texts),
                    'file_type': 'pdf'
                },
                'page_texts': page_texts  # Useful for page-specific retrieval
            }
            
        except Exception as e:
            error_msg = f"PDF processing error: {str(e)}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            return {
                'success': False,
                'error': error_msg
            }
    
    # ========================================================================
    # WORD DOCUMENT PROCESSING
    # ========================================================================
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word documents.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        print("[DocumentProcessor] Extracting text from Word document...")
        
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'Word support not available. Install python-docx.'
            }
        
        try:
            # Open the Word document
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ""
            word_count = 0
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                
                if para_text:  # Skip empty paragraphs
                    paragraphs.append({
                        'paragraph': para_num,
                        'text': para_text,
                        'style': paragraph.style.name if paragraph.style else None
                    })
                    
                    full_text += para_text + "\n\n"
                    word_count += len(para_text.split())
            
            # Also extract text from tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                table_text = "\n[Table]\n"
                
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text += row_text + "\n"
                
                full_text += table_text + "\n"
                word_count += len(table_text.split())
            
            print(f"[DocumentProcessor] ✅ Extracted {len(paragraphs)} paragraphs")
            print(f"[DocumentProcessor] ✅ Extracted {table_count} tables")
            print(f"[DocumentProcessor] Total words: {word_count:,}")
            
            return {
                'success': True,
                'text': full_text.strip(),
                'metadata': {
                    'paragraph_count': len(paragraphs),
                    'table_count': table_count,
                    'word_count': word_count,
                    'file_type': 'docx'
                },
                'paragraphs': paragraphs  # Useful for structure-aware retrieval
            }
            
        except Exception as e:
            error_msg = f"Word processing error: {str(e)}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            return {
                'success': False,
                'error': error_msg
            }
    
    # ========================================================================
    # EXCEL PROCESSING
    # ========================================================================
    
    def process_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Excel files.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        print("[DocumentProcessor] Extracting text from Excel file...")
        
        if not EXCEL_AVAILABLE:
            return {
                'success': False,
                'error': 'Excel support not available. Install openpyxl.'
            }
        
        try:
            # Open the Excel workbook
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            sheet_names = workbook.sheetnames
            print(f"[DocumentProcessor] Excel has {len(sheet_names)} sheets: {sheet_names}")
            
            full_text = ""
            sheet_data = []
            total_cells = 0
            
            # Process each sheet
            for sheet_name in sheet_names:
                sheet = workbook[sheet_name]
                print(f"  Processing sheet: {sheet_name}")
                
                sheet_text = f"\n[Sheet: {sheet_name}]\n"
                row_count = 0
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    # Filter out completely empty rows
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(
                            str(cell) if cell is not None else ""
                            for cell in row
                        )
                        sheet_text += row_text + "\n"
                        row_count += 1
                        total_cells += sum(1 for cell in row if cell is not None)
                
                if row_count > 0:
                    sheet_data.append({
                        'sheet': sheet_name,
                        'rows': row_count,
                        'text': sheet_text
                    })
                    full_text += sheet_text + "\n"
                    
                print(f"    - Extracted {row_count} rows")
            
            # Close the workbook
            workbook.close()
            
            word_count = len(full_text.split())
            print(f"[DocumentProcessor] ✅ Extracted data from {len(sheet_data)} sheets")
            print(f"[DocumentProcessor] Total cells with data: {total_cells:,}")
            print(f"[DocumentProcessor] Total words: {word_count:,}")
            
            return {
                'success': True,
                'text': full_text.strip(),
                'metadata': {
                    'sheet_count': len(sheet_names),
                    'sheets_with_data': len(sheet_data),
                    'total_cells': total_cells,
                    'word_count': word_count,
                    'file_type': 'excel'
                },
                'sheet_data': sheet_data  # Useful for sheet-specific queries
            }
            
        except Exception as e:
            error_msg = f"Excel processing error: {str(e)}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            return {
                'success': False,
                'error': error_msg
            }
    
    # ========================================================================
    # PLAIN TEXT PROCESSING
    # ========================================================================
    
    def process_text(self, file_path: str) -> Dict[str, Any]:
        """
        Process plain text files (.txt).
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dictionary with text and metadata
        """
        print("[DocumentProcessor] Processing plain text file...")
        
        try:
            # Read the text file
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            # Calculate metadata
            line_count = len(text.splitlines())
            word_count = len(text.split())
            char_count = len(text)
            
            print(f"[DocumentProcessor] ✅ Read text file")
            print(f"[DocumentProcessor] Lines: {line_count:,}")
            print(f"[DocumentProcessor] Words: {word_count:,}")
            print(f"[DocumentProcessor] Characters: {char_count:,}")
            
            return {
                'success': True,
                'text': text,
                'metadata': {
                    'line_count': line_count,
                    'word_count': word_count,
                    'character_count': char_count,
                    'file_type': 'txt'
                }
            }
            
        except Exception as e:
            error_msg = f"Text file processing error: {str(e)}"
            print(f"[DocumentProcessor] ❌ {error_msg}")
            return {
                'success': False,
                'error': error_msg
            }
    
    # ========================================================================
    # TEXT CHUNKING
    # ========================================================================
    
    def create_chunks(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into chunks for vector embedding.
        
        Why chunking?
        - Language models have token limits
        - Smaller chunks = more precise retrieval
        - Overlapping chunks preserve context
        
        Args:
            text: Full text to chunk
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        print(f"\n[DocumentProcessor] Creating chunks...")
        print(f"  Text length: {len(text):,} characters")
        
        if not text or len(text.strip()) == 0:
            print("[DocumentProcessor] ⚠️  No text to chunk")
            return []
        
        chunks = []
        
        if LANGCHAIN_AVAILABLE and self.text_splitter:
            # Use LangChain's advanced splitter
            print("[DocumentProcessor] Using LangChain RecursiveCharacterTextSplitter")
            
            # Split the text
            text_chunks = self.text_splitter.split_text(text)
            
            # Create chunk dictionaries with metadata
            for i, chunk_text in enumerate(text_chunks):
                chunk_data = {
                    'chunk_index': i,
                    'text': chunk_text.strip(),
                    'char_count': len(chunk_text),
                    'word_count': len(chunk_text.split()),
                    'start_char': text.find(chunk_text[:50]) if len(chunk_text) >= 50 else 0
                }
                chunks.append(chunk_data)
                
                # Debug: Show first few chunks
                if i < 3:
                    preview = chunk_text[:100] + "..." if len(chunk_text) > 100 else chunk_text
                    print(f"  Chunk {i}: {preview}")
        
        else:
            # Basic chunking fallback
            print("[DocumentProcessor] Using basic chunking (no LangChain)")
            
            # Simple splitting by character count
            chunk_size = self.chunk_size
            overlap = self.chunk_overlap
            
            start = 0
            chunk_index = 0
            
            while start < len(text):
                # Calculate end position
                end = start + chunk_size
                
                # Try to break at a sentence boundary
                if end < len(text):
                    # Look for sentence end (.!?) near the chunk boundary
                    for punct in ['. ', '! ', '? ', '\n\n', '\n']:
                        last_punct = text.rfind(punct, start, end)
                        if last_punct != -1 and last_punct > start + chunk_size // 2:
                            end = last_punct + len(punct)
                            break
                
                # Extract chunk
                chunk_text = text[start:end].strip()
                
                if chunk_text:  # Only add non-empty chunks
                    chunk_data = {
                        'chunk_index': chunk_index,
                        'text': chunk_text,
                        'char_count': len(chunk_text),
                        'word_count': len(chunk_text.split()),
                        'start_char': start
                    }
                    chunks.append(chunk_data)
                    chunk_index += 1
                    
                    # Debug: Show first few chunks
                    if chunk_index <= 3:
                        preview = chunk_text[:100] + "..." if len(chunk_text) > 100 else chunk_text
                        print(f"  Chunk {chunk_index-1}: {preview}")
                
                # Move to next chunk with overlap
                start = end - overlap if end < len(text) else end
        
        print(f"[DocumentProcessor] ✅ Created {len(chunks)} chunks")
        
        # Show chunk statistics
        if chunks:
            avg_size = sum(c['char_count'] for c in chunks) / len(chunks)
            min_size = min(c['char_count'] for c in chunks)
            max_size = max(c['char_count'] for c in chunks)
            
            print(f"[DocumentProcessor] Chunk statistics:")
            print(f"  - Average size: {avg_size:.0f} characters")
            print(f"  - Min size: {min_size} characters")
            print(f"  - Max size: {max_size} characters")
        
        return chunks

# ============================================================================
# STANDALONE TEST FUNCTION
# ============================================================================

def test_document_processor():
    """
    Test the document processor with sample files.
    Run this to verify everything works.
    """
    print("\n" + "="*70)
    print("DOCUMENT PROCESSOR TEST")
    print("="*70)
    
    # Initialize processor
    processor = DocumentProcessor()
    
    # Create test files
    test_files = []
    
    # Test 1: Create a test text file
    test_txt = Path("test_document.txt")
    test_content = """Introduction to RAG Systems

Retrieval-Augmented Generation (RAG) is a powerful technique that combines 
the strengths of large language models with external knowledge retrieval.

Key Components:
1. Document Processing: Extract and chunk text from various sources
2. Embedding Generation: Convert text chunks into vector representations
3. Vector Storage: Store embeddings in a vector database like Pinecone
4. Retrieval: Find relevant chunks based on query similarity
5. Generation: Use retrieved context to generate accurate responses

Benefits of RAG:
- Reduced hallucination
- Access to private/updated information
- Traceable sources
- Cost-effective compared to fine-tuning

This test document demonstrates the document processor's ability to handle
text extraction and chunking for the RAG pipeline.
"""
    
    with open(test_txt, 'w', encoding='utf-8') as f:
        f.write(test_content)
    test_files.append(test_txt)
    
    print(f"\nCreated test file: {test_txt}")
    print(f"Size: {len(test_content)} characters")
    
    # Process the test file
    print(f"\n{'='*50}")
    print("Testing document processing...")
    print('='*50)
    
    result = processor.process_document(str(test_txt))
    
    # Display results
    if result['success']:
        print(f"\n✅ Processing successful!")
        print(f"\nMetadata:")
        for key, value in result.get('metadata', {}).items():
            print(f"  - {key}: {value}")
        
        print(f"\nChunks created: {len(result.get('chunks', []))}")
        
        # Show first chunk as example
        if result.get('chunks'):
            first_chunk = result['chunks'][0]
            print(f"\nFirst chunk example:")
            print(f"  Index: {first_chunk['chunk_index']}")
            print(f"  Size: {first_chunk['char_count']} characters")
            print(f"  Words: {first_chunk['word_count']}")
            print(f"  Text preview: {first_chunk['text'][:200]}...")
    else:
        print(f"\n❌ Processing failed: {result.get('error')}")
    
    # Clean up test files
    print(f"\n{'='*50}")
    print("Cleaning up test files...")
    for test_file in test_files:
        if test_file.exists():
            test_file.unlink()
            print(f"  Deleted: {test_file}")
    
    print("\n" + "="*70)
    print("TEST COMPLETE")
    print("="*70)

# Run test if this file is executed directly
if __name__ == "__main__":
    test_document_processor()